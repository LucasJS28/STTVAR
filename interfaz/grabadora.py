import sys
import os
import json
import time
import queue
import wave
import socket
import threading
from datetime import datetime
from io import StringIO, BytesIO

# --- Dependencias de Terceros ---
import sounddevice as sd
import numpy as np
import argostranslate.package
import argostranslate.translate
import pyttsx3
import requests
from docx import Document
from flask import Flask, Response, render_template, request, jsonify, send_from_directory
from pyngrok import ngrok
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout,
                             QPushButton, QLabel, QComboBox, QStackedWidget,
                             QColorDialog, QSlider, QFontComboBox, QFrame, QGraphicsDropShadowEffect,
                             QGridLayout, QSizePolicy, QGraphicsOpacityEffect, QMessageBox,
                             QProgressDialog, QInputDialog)
from PyQt5.QtGui import QIcon, QColor, QFont, QPixmap, QCursor
from PyQt5.QtCore import (Qt, QTimer, QPoint, QPropertyAnimation, QEasingCurve,
                          QRect, QByteArray, QSize, pyqtSignal)

# --- Dependencias del Proyecto (Asegúrate de que estos archivos existan) ---
try:
    from transcripcion.transcriber import TranscriberThread
    from transcripcion.vosk_utils import SAMPLE_RATE
    from interfaz.menu import NuevaVentana
    from interfaz.configuraciones import ConfiguracionIA
except ImportError as e:
    print(f"Error de importación: {e}")
    # Esta sección se deja para depuración, pero no debería fallar si la estructura es correcta.


# =====================================================================
# Sección de Iconos SVG
# =====================================================================
def create_icon_from_svg(svg_data_str: str, color: str) -> QIcon:
    """Crea un QIcon a partir de datos SVG, reemplazando el color de relleno."""
    final_svg = svg_data_str.replace('fill="white"', f'fill="{color}"')
    byte_array = QByteArray(final_svg.encode('utf-8'))
    pixmap = QPixmap()
    pixmap.loadFromData(byte_array)
    return QIcon(pixmap)

icon_data = {
    "mic": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M12 14c1.66 0 3-1.34 3-3V5c0-1.66-1.34-3-3-3S9 3.34 9 5v6c0 1.66 1.34 3 3 3zm5.3-3c0 3-2.54 5.1-5.3 5.1S6.7 14 6.7 11H5c0 3.41 2.72 6.23 6 6.72V21h2v-3.28c3.28-.49 6-3.31 6-6.72h-1.7z"/></svg>',
    "mic_off": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M19 11h-1.7c0 .58-.1 1.13-.27 1.64l1.27 1.27c.44-.88.7-1.87.7-2.91zm-6 5.1c1.13 0 2.16-.39 3-1.02l-1.46-1.46c-.47.29-1.02.48-1.54.48V14c1.66 0 3-1.34 3-3V5c0-1.06-.55-2-1.33-2.54l-2.09 2.09c.27.16.52.36.72.6V11c0 .19-.03.37-.08.54l-1.41-1.41C11.89 9.58 12 8.8 12 8V7h-1.7l-2-2H12c1.66 0 3 1.34 3 3v.18l-3-3V5c0-1.66-1.34-3-3-3s-3 1.34-3 3v1.18l-1.4-1.4C5.55 5 6.7 4 8 4c.48 0 .93.11 1.33.3L1.39 4.22 2.8 5.63 9 11.83V14c0 1.66 1.34 3 3 3v2.28c-3.28-.49-6-3.31-6-6.72H5c0 3.41 2.72 6.23 6 6.72V22h2v-1.28c.45-.07.88-.2 1.3-.38l.64.64 1.41-1.41-11.7-11.7z"/></svg>',
    "play": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M8 5v14l11-7z"/></svg>',
    "unlock": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M12 17c1.1 0 2-.9 2-2s-.9-2-2-2-2 .9-2 2 .9 2 2 2zm6-9h-1V6c0-2.76-2.24-5-5-5S7 3.24 7 6h1.9c0-1.71 1.39-3.1 3.1-3.1 1.71 0 3.1 1.39 3.1 3.1v2H6c-1.1 0-2 .9-2 2v10c0 1.1.9 2 2 2h12c1.1 0 2-.9 2-2V10c0-1.1-.9-2-2-2zm0 12H6V10h12v10z"/></svg>',
    "lock": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M18 8h-1V6c0-2.76-2.24-5-5-5S7 3.24 7 6v2H6c-1.1 0-2 .9-2 2v10c0 1.1.9 2 2 2h12c1.1 0 2-.9 2-2V10c0-1.1-.9-2-2-2zm-6 9c-1.1 0-2-.9-2-2s.9-2 2-2 2 .9 2 2-.9 2-2 2zm3.1-9H8.9V6c0-1.71 1.39-3.1 3.1-3.1 1.71 0 3.1 1.39 3.1 3.1v2z"/></svg>',
    "record_active": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10 10-4.48 10-10S17.52 2 12 2z"/></svg>',
    "stop": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M6 6h12v12H6z"/></svg>',
    "palette": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M12 3c-4.97 0-9 4.03-9 9s4.03 9 9 9c.83 0 1.5-.67 1.5-1.5 0-.39-.15-.74-.39-1.01-.23-.26-.38-.61-.38-.99 0-.83.67-1.5 1.5-1.5H16c2.76 0 5-2.24 5-5 0-4.42-4.03-8-9-8zm-5.5 9c-.83 0-1.5-.67-1.5-1.5S5.67 9 6.5 9 8 9.67 8 10.5 7.33 12 6.5 12zm3-4C8.67 8 8 7.33 8 6.5S8.67 5 9.5 5s1.5.67 1.5 1.5S10.33 8 9.5 8zm5 0c-.83 0-1.5-.67-1.5-1.5S13.67 5 14.5 5s1.5.67 1.5 1.5S15.33 8 14.5 8zm3 4c-.83 0-1.5-.67-1.5-1.5S16.67 9 17.5 9s1.5.67 1.5 1.5-.67 1.5-1.5 1.5z"/></svg>',
    "font_color": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M9.93 13.5h4.14L12 7.98zM20 2H4c-1.1 0-2 .9-2 2v16c0 1.1.9 2 2 2h16c1.1 0 2-.9 2-2V4c0-1.1-.9-2-2-2zm-4.05 16.5-1.14-3H9.17l-1.12 3H5.96l5.11-13h1.86l5.11 13h-2.09z"/></svg>',
    "transparency": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10 10-4.48 10-10S17.52 2 12 2zM12 18c-3.31 0-6-2.69-6-6s2.69-6 6-6v12z"/></svg>',
    "minus": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M5 11h14v2H5z"/></svg>',
    "plus": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M19 11h-6V5h-2v6H5v2h6v6h2v-6h6z"/></svg>',
    "menu": '<svg xmlns="http://www.w3.org/2000/svg" height="24" viewBox="0 96 960 960" width="24" fill="white"><path d="M120 816v-60h720v60H120Zm0-210v-60h720v60H120Zm0-210v-60h720v60H120Z"/></svg>',
    "font_family": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M5 4v3h5.5v12h3V7H19V4z"/></svg>',
    "language": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M11.99 2C6.47 2 2 6.48 2 12s4.47 10 9.99 10C17.52 22 22 17.52 22 12S17.52 2 11.99 2zM12 20c-4.42 0-8-3.58-8-8s3.58-8 8-8 8 3.58 8 8-3.58 8-8 8zm.5-13H11v6l5.25 3.15.75-1.23-4.5-2.67V7z"/></svg>',
    "close": '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="white"><path d="M19 6.41L17.59 5 12 10.59 6.41 5 5 6.41 10.59 12 5 17.59 6.41 19 12 13.41 17.59 19 19 17.59 13.41 12z"/></svg>',
    "share": '<svg xmlns="http://www.w3.org/2000/svg" height="24" viewBox="0 0 24 24" width="24" fill="white"><path d="M18 16.08c-.76 0-1.44.3-1.96.77L8.91 12.7c.05-.23.09-.46.09-.7s-.04-.47-.09-.7l7.05-4.11c.54.5 1.25.81 2.04.81 1.66 0 3-1.34 3-3s-1.34-3-3-3-3 1.34-3 3c0 .24.04.47.09.7L8.04 8.41C7.5 7.91 6.79 7.6 6 7.6c-1.66 0-3 1.34-3 3s1.34 3 3 3c.79 0 1.5-.31 2.04-.81l7.12 4.16c-.05.21-.08.43-.08.65 0 1.61 1.31 2.92 2.92 2.92s2.92-1.31 2.92-2.92-1.31-2.92-2.92-2.92z"/></svg>',
    "settings": '<svg xmlns="http://www.w3.org/2000/svg" height="24" viewBox="0 0 24 24" width="24" fill="white"><path d="M19.43 12.98c.04-.32.07-.64.07-.98s-.03-.66-.07-.98l2.11-1.65c.19-.15.24-.42.12-.64l-2-3.46c-.12-.22-.39-.3-.61-.22l-2.49 1c-.52-.4-1.08-.73-1.69-.98l-.38-2.65C14.46 2.18 14.25 2 14 2h-4c-.25 0-.46.18-.49.42l-.38 2.65c-.61.25-1.17.59-1.69-.98l-2.49-1c-.23-.09-.49 0-.61.22l-2 3.46c-.13.22-.07.49.12.64l2.11 1.65c-.04.32-.07.65-.07.98s.03.66.07.98l-2.11 1.65c-.19.15-.24.42-.12-.64l2 3.46c.12.22.39.3.61.22l2.49-1c.52.4 1.08.73 1.69.98l.38 2.65c.03.24.24.42.49.42h4c.25 0 .46-.18.49-.42l.38-2.65c.61-.25 1.17-.59 1.69-.98l2.49 1c.23.09.49 0 .61.22l2-3.46c.12-.22-.07-.49-.12-.64l-2.11-1.65zM12 15.5c-1.93 0-3.5-1.57-3.5-3.5s1.57-3.5 3.5-3.5 3.5 1.57 3.5 3.5-1.57 3.5-3.5 3.5z"/></svg>',
    "files": '<svg xmlns="http://www.w3.org/2000/svg" height="24" viewBox="0 0 24 24" width="24" fill="white"><path d="M10 4H4c-1.1 0-1.99.9-1.99 2L2 18c0 1.1.9 2 2 2h16c1.1 0 2-.9 2-2V8c0-1.1-.9-2-2-2h-8l-2-2z"/></svg>',
}


# =====================================================================
# Sección de Flask y Variables Globales
# =====================================================================
# MODIFICADO: Busca 'templates' relativo a la ubicación de este script
template_folder_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
app = Flask(__name__, template_folder=template_folder_path)
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0

transcription_chunks = []
transcripcion_con_timestamps = ""
transcripcion_plana = ""
current_partial = ""
client_ai_responses = {}
clientes_idioma = {}
clientes_idioma_acumulado = {}
chat_messages = []
qtextedit_buffer = ""
last_qtextedit_update = 0
qtextedit_update_interval = 0.5
shared_audio_data = {}
active_clients = {}
shared_in_memory_audio_buffer = []

model = None
q = queue.Queue(maxsize=20)
OLLAMA_API_URL = "http://localhost:11434/api/generate"
CONFIG_FILE = 'config.json'
MODEL_NAME = None

OLLAMA_AVAILABLE = False
OLLAMA_STATUS_MESSAGE = "Iniciando..."

# =====================================================================
# Sesión HTTP y Funciones de Ayuda
# =====================================================================
http_session = requests.Session()

def load_config():
    global MODEL_NAME
    default_model = "phi3.5:latest"
    try:
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                MODEL_NAME = json.load(f).get('ollama_model', default_model)
        else:
            MODEL_NAME = default_model
    except (IOError, json.JSONDecodeError):
        MODEL_NAME = default_model
    print(f"Modelo de IA por defecto cargado: {MODEL_NAME}")

def update_ollama_status():
    global OLLAMA_AVAILABLE, OLLAMA_STATUS_MESSAGE
    load_config()
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=3)
        if response.status_code == 200:
            models = response.json().get("models", [])
            if not models:
                OLLAMA_AVAILABLE = False
                OLLAMA_STATUS_MESSAGE = "Error: IA activa, pero sin modelos. Usa 'ollama pull <modelo>'."
            elif any(m['name'] == MODEL_NAME for m in models):
                OLLAMA_AVAILABLE = True
                OLLAMA_STATUS_MESSAGE = f"Ollama conectado con el modelo {MODEL_NAME}."
            else:
                OLLAMA_AVAILABLE = False
                OLLAMA_STATUS_MESSAGE = f"Error: Modelo '{MODEL_NAME}' no instalado. Selecciona otro en la config 🧠."
        else:
            OLLAMA_AVAILABLE = False
            OLLAMA_STATUS_MESSAGE = f"Error: No se pudo comunicar con Ollama (Código: {response.status_code})."
    except requests.exceptions.RequestException:
        OLLAMA_AVAILABLE = False
        OLLAMA_STATUS_MESSAGE = "Error: IA no disponible. Asegúrate de que Ollama esté en ejecución."
    print(f"Estado de Ollama: {OLLAMA_STATUS_MESSAGE}")

update_ollama_status()

def setup_argos():
    # << CORRECCIÓN: Lógica para verificar paquetes instalados
    try:
        argostranslate.package.update_package_index()
        available_packages = argostranslate.package.get_available_packages()
        installed_packages = argostranslate.package.get_installed_packages()
        
        installed_lang_codes = set((p.from_code, p.to_code) for p in installed_packages)

        for lang_to in ["en", "fr", "de"]:
            if ("es", lang_to) not in installed_lang_codes:
                package_to_install = next((p for p in available_packages if p.from_code == "es" and p.to_code == lang_to), None)
                if package_to_install:
                    print(f"Instalando paquete de traducción es -> {lang_to}")
                    package_to_install.install()
    except Exception as e:
        print(f"Error al instalar paquetes de traducción: {e}")

setup_argos()
installed_languages = argostranslate.translate.get_installed_languages()
source_lang = next((l for l in installed_languages if l.code == "es"), None)


# =====================================================================
# Rutas de la Aplicación Flask (Sin cambios en esta sección)
# =====================================================================
@app.route("/")
def index():
    return render_template('index.html', ollama_is_available=OLLAMA_AVAILABLE)

@app.route('/chat/send', methods=['POST'])
def send_chat_message():
    data = request.json
    if not data or 'user' not in data: return jsonify({"status": "error", "message": "Falta el usuario"}), 400
    is_note = data.get('isNote', False)
    if is_note:
        if 'content' not in data or not data['content'].strip(): return jsonify({"status": "error", "message": "Los apuntes están vacíos"}), 400
        message_data = {"user": data['user'], "isNote": True, "content": data['content'], "timestamp": time.time()}
    else:
        if 'message' not in data or not data['message'].strip(): return jsonify({"status": "error", "message": "El mensaje está vacío"}), 400
        message_data = {"user": data['user'], "isNote": False, "message": data['message'], "timestamp": time.time()}
    chat_messages.append(message_data)
    if len(chat_messages) > 100: chat_messages.pop(0)
    return jsonify({"status": "ok"})

@app.route("/stream")
def stream():
    client_id = request.args.get("client_id", "default")
    
    def event_stream():
        global active_clients
        try:
            active_clients[client_id] = time.time()
            print(f"Cliente conectado: {client_id}. Total de usuarios: {len(active_clients)}")

            last_chat_message_count = len(chat_messages)
            last_sent_data = {}
            last_sent_user_count = 0

            while True:
                current_user_count = len(active_clients)
                if current_user_count != last_sent_user_count:
                    system_payload = {"type": "system_info", "data": {"user_count": current_user_count}}
                    yield f"data: {json.dumps(system_payload)}\n\n"
                    last_sent_user_count = current_user_count
                
                idioma_sub = clientes_idioma.get(client_id, "es")
                idioma_acum = clientes_idioma_acumulado.get(client_id, "es")
                
                subtitulo_text = current_partial
                acumulado_text_flat = transcripcion_plana
                acumulado_chunks_final = transcription_chunks
                ai_response_text = client_ai_responses.get(client_id, "")

                if source_lang:
                    if idioma_sub != "es" and subtitulo_text.strip():
                        try:
                            target_lang = next((l for l in installed_languages if l.code == idioma_sub), None)
                            if target_lang: subtitulo_text = source_lang.get_translation(target_lang).translate(subtitulo_text)
                        except Exception as e: print(f"Error traduciendo subtítulo: {e}")
                    
                    if idioma_acum != "es":
                        try:
                            target_lang_acum = next((l for l in installed_languages if l.code == idioma_acum), None)
                            if target_lang_acum:
                                if acumulado_text_flat.strip():
                                    acumulado_text_flat = source_lang.get_translation(target_lang_acum).translate(acumulado_text_flat)
                                
                                translated_chunks = []
                                for chunk in transcription_chunks:
                                    translated_text = source_lang.get_translation(target_lang_acum).translate(chunk["text"])
                                    new_chunk = chunk.copy()
                                    new_chunk["text"] = translated_text
                                    translated_chunks.append(new_chunk)
                                acumulado_chunks_final = translated_chunks
                        except Exception as e: print(f"Error traduciendo acumulado: {e}")

                current_data = {
                    "subtitulo": subtitulo_text,
                    "acumulado_chunks": acumulado_chunks_final,
                    "acumulado_flat": acumulado_text_flat,
                    "ai_response": ai_response_text
                }
                
                if current_data != last_sent_data.get("transcript"):
                    last_sent_data["transcript"] = current_data
                    transcript_payload = {"type": "transcript", "data": current_data}
                    yield f"data: {json.dumps(transcript_payload)}\n\n"
                
                if len(chat_messages) > last_chat_message_count:
                    new_messages = chat_messages[last_chat_message_count:]
                    for msg in new_messages: yield f"data: {json.dumps({'type': 'chat', 'data': msg})}\n\n"
                    last_chat_message_count = len(chat_messages)
                
                time.sleep(0.1)
        finally:
            active_clients.pop(client_id, None)
            print(f"Cliente desconectado: {client_id}. Total de usuarios: {len(active_clients)}")

    return Response(event_stream(), mimetype="text/event-stream")

@app.route('/get_audio_chunk')
def get_audio_chunk():
    global shared_in_memory_audio_buffer
    start_sec_str = request.args.get('start')
    end_sec_str = request.args.get('end')

    if not all([start_sec_str, end_sec_str]): return "Faltan parámetros de tiempo.", 400
    if not shared_in_memory_audio_buffer: return "No hay datos de audio en memoria.", 404

    try:
        start_sec, end_sec = float(start_sec_str), float(end_sec_str)
        sample_width, channels, frame_rate = 2, 1, SAMPLE_RATE
        bytes_per_second = frame_rate * sample_width * channels
        start_byte, end_byte = int(start_sec * bytes_per_second), int(end_sec * bytes_per_second)
        
        full_audio_data = b''.join(shared_in_memory_audio_buffer)
        audio_chunk_data = full_audio_data[start_byte:end_byte]

        buffer = BytesIO()
        with wave.open(buffer, 'wb') as wf:
            wf.setnchannels(channels)
            wf.setsampwidth(sample_width)
            wf.setframerate(frame_rate)
            wf.writeframes(audio_chunk_data)
        buffer.seek(0)
        
        return Response(buffer, mimetype="audio/wav", headers={"Content-Disposition": "inline"})
    except Exception as e:
        print(f"Error al procesar chunk de audio desde memoria (método manual): {e}")
        return "Error interno del servidor al procesar el audio.", 500

@app.route("/set_language/<client_id>/<lang>", methods=['POST'])
def set_language(client_id, lang): clientes_idioma[client_id] = lang; return "OK"

@app.route("/set_language_acumulado/<client_id>/<lang>", methods=['POST'])
def set_language_acumulado(client_id, lang): clientes_idioma_acumulado[client_id] = lang; return "OK"

@app.route("/clear_text/<client_id>", methods=['POST'])
def clear_text(client_id):
    global transcripcion_con_timestamps, transcripcion_plana, current_partial, qtextedit_buffer, shared_audio_data, transcription_chunks, shared_in_memory_audio_buffer
    transcripcion_con_timestamps = ""; transcripcion_plana = ""; current_partial = ""; qtextedit_buffer = ""
    transcription_chunks = []; shared_in_memory_audio_buffer = []
    client_ai_responses.pop(client_id, None); shared_audio_data.clear()
    return "OK"

@app.route("/ask_ai/<client_id>", methods=['POST'])
def ask_ai(client_id):
    data = request.get_json()
    if not data: return "Bad Request: No JSON data received", 400
    question = data.get("question", "")
    ai_thread = threading.Thread(target=process_ai_question, args=(question, client_id)); ai_thread.start()
    return "OK"

@app.route("/download_txt/<client_id>/<filename>", methods=['POST'])
def download_txt(client_id, filename):
    acumulado_text = transcripcion_con_timestamps
    buffer = StringIO(); buffer.write(f"Transcripción de STTTVAR\nFecha: {datetime.now().strftime('%d/%m/%Y, %H:%M:%S')}\n\n{acumulado_text}")
    txt_content = buffer.getvalue(); buffer.close()
    return Response(txt_content, mimetype='text/plain', headers={'Content-Disposition': f'attachment; filename={filename}.txt'})

@app.route("/download_docx/<client_id>/<filename>", methods=['POST'])
def download_docx(client_id, filename):
    acumulado_text = transcripcion_con_timestamps
    doc = Document(); doc.add_heading("Transcripción de STTTVAR", 0); doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d de %B de %Y, %I:%M %p')}"); doc.add_heading("Transcripción:", level=1); doc.add_paragraph(acumulado_text)
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return Response(buffer.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': f'attachment; filename={filename}.docx'})

@app.route('/audio_status')
def audio_status():
    return jsonify({'available': bool(shared_audio_data and shared_audio_data.get('available')), 'filename': shared_audio_data.get('filename')})

@app.route('/download_audio')
def download_audio():
    if shared_audio_data and shared_audio_data.get('available'):
        try: return send_from_directory(shared_audio_data['directory'], shared_audio_data['filename'], as_attachment=True)
        except Exception as e: print(f"Error al servir el archivo de audio: {e}"); return "Error al leer el archivo.", 500
    return "No hay archivo de audio disponible para descargar.", 404

def process_ai_question(question, client_id):
    if not OLLAMA_AVAILABLE:
        client_ai_responses[client_id] = OLLAMA_STATUS_MESSAGE
        return
    if not question.strip():
        client_ai_responses[client_id] = "Error: La pregunta no puede estar vacía."
        return
    try:
        context = transcripcion_plana.strip() or "No hay contexto disponible."
        if len(context.split()) > 1000:
            context = " ".join(context.split()[-1000:])
        prompt = (f"Usando el siguiente texto como referencia:\n\n{context}\n\nPregunta: {question}\nRespuesta concisa en español:")
        payload = {"model": MODEL_NAME, "prompt": prompt, "stream": True, "options": {"temperature": 0.6, "num_predict": 256}}
        client_ai_responses[client_id] = ""
        with http_session.post(OLLAMA_API_URL, json=payload, stream=True, timeout=60) as response:
            if response.status_code == 200:
                for line in response.iter_lines():
                    if line:
                        try:
                            chunk = json.loads(line)
                            content = chunk.get("response", "")
                            if content: client_ai_responses[client_id] += content
                            if chunk.get("done"): break
                        except json.JSONDecodeError: print(f"Error decodificando la respuesta JSON de Ollama: {line}")
            else:
                client_ai_responses[client_id] = f"Error en la API de Ollama: Código {response.status_code}."
    except requests.exceptions.RequestException as e:
        client_ai_responses[client_id] = f"Error de conexión con Ollama: {e}"
    except Exception as e:
        client_ai_responses[client_id] = f"Error inesperado al procesar la pregunta: {e}"

# =====================================================================
# Clases de la Interfaz Gráfica (PyQt5)
# =====================================================================

class OptionsMenu(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.Popup | Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.main_layout = QHBoxLayout(self)
        self.main_layout.setContentsMargins(10, 10, 10, 10)
        self.main_layout.setSpacing(6)
        
        self.background_widget = QFrame(self)
        self.background_widget.setObjectName("optionsMenuBackground")
        
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(20)
        shadow.setXOffset(0)
        shadow.setYOffset(4)
        shadow.setColor(QColor(0, 0, 0, 160))
        self.background_widget.setGraphicsEffect(shadow)

        layout_container = QHBoxLayout(self.background_widget)
        layout_container.setContentsMargins(5, 5, 5, 5)
        layout_container.setSpacing(6)
        self.main_layout.addWidget(self.background_widget)
        
        self.lock_btn = self._create_menu_button(icon_data["unlock"], "Bloquear posición")
        self.lock_btn.setCheckable(True)
        self.bg_color_btn = self._create_menu_button(icon_data["palette"], "Color de fondo")
        self.font_color_btn = self._create_menu_button(icon_data["font_color"], "Color de fuente")
        self.language_selector_wrapper = self._create_compound_widget()
        lang_layout = self.language_selector_wrapper.layout()
        lang_icon = QLabel()
        lang_icon.setPixmap(create_icon_from_svg(icon_data["language"], "white").pixmap(16, 16))
        self.language_selector = QComboBox()
        self.language_selector.addItems(["Español", "Inglés", "Portugués"])
        self.language_selector.setFixedWidth(80)
        lang_layout.addWidget(lang_icon)
        lang_layout.addWidget(self.language_selector)
        
        opacity_widget = self._create_compound_widget()
        opacity_layout = opacity_widget.layout()
        opacity_icon = QLabel()
        opacity_icon.setPixmap(create_icon_from_svg(icon_data["transparency"], "white").pixmap(16, 16))
        self.opacity_slider = QSlider(Qt.Horizontal)
        self.opacity_slider.setRange(10, 100)
        self.opacity_slider.setValue(85)
        self.opacity_slider.setFixedWidth(50)
        opacity_layout.addWidget(opacity_icon)
        opacity_layout.addWidget(self.opacity_slider)
        
        self.font_selector_wrapper = self._create_compound_widget()
        font_layout = self.font_selector_wrapper.layout()
        font_icon = QLabel()
        font_icon.setPixmap(create_icon_from_svg(icon_data["font_family"], "white").pixmap(16, 16))
        self.font_selector = QFontComboBox()
        self.font_selector.setEditable(False)
        self.font_selector.setFixedWidth(120)
        font_layout.addWidget(font_icon)
        font_layout.addWidget(self.font_selector)

        self.decrease_font_btn = self._create_menu_button(icon_data["minus"], "Disminuir fuente")
        self.increase_font_btn = self._create_menu_button(icon_data["plus"], "Aumentar fuente")
        
        self.rec_btn = self._create_menu_button(icon_data["record_active"], "Iniciar grabación")
        self.rec_btn.setObjectName("rec_btn")
        self.rec_btn.setCheckable(True)
        
        self.mute_btn = self._create_menu_button(icon_data["mic"], "Silenciar micrófono")
        self.mute_btn.setObjectName("mic_btn")
        self.mute_btn.setCheckable(True)

        self.share_btn = self._create_menu_button(icon_data["share"], "Compartir sesión en web")
        self.share_btn.setObjectName("share_btn")
        
        layout_container.addWidget(self.lock_btn)
        layout_container.addWidget(self._create_divider())
        layout_container.addWidget(self.bg_color_btn)
        layout_container.addWidget(self.font_color_btn)
        layout_container.addWidget(self.language_selector_wrapper)
        layout_container.addWidget(opacity_widget)
        layout_container.addWidget(self.font_selector_wrapper)
        layout_container.addWidget(self.decrease_font_btn)
        layout_container.addWidget(self.increase_font_btn)
        
        layout_container.addStretch(1)
        
        layout_container.addWidget(self._create_divider())
        layout_container.addWidget(self.rec_btn)
        layout_container.addWidget(self.mute_btn)
        layout_container.addWidget(self.share_btn)
        
        self.setLayout(self.main_layout)
        
        self.fade_animation = QPropertyAnimation(self, b"windowOpacity")
        self.fade_animation.setDuration(200)
        self.fade_animation.setEasingCurve(QEasingCurve.OutCubic)

        self.rec_opacity_effect = QGraphicsOpacityEffect(self.rec_btn)
        self.rec_btn.setGraphicsEffect(self.rec_opacity_effect)
        self.rec_animation = QPropertyAnimation(self.rec_opacity_effect, b"opacity")
        self.rec_animation.setDuration(800)
        self.rec_animation.setStartValue(1.0)
        self.rec_animation.setEndValue(0.4)
        self.rec_animation.setLoopCount(-1)
        self.rec_animation.setEasingCurve(QEasingCurve.InOutCubic)

    def _create_menu_button(self, svg_data, tooltip):
        btn = QPushButton()
        btn.setIcon(create_icon_from_svg(svg_data, "#f0f0f0"))
        btn.setIconSize(QSize(16, 16))
        btn.setToolTip(tooltip)
        btn.setFixedSize(32, 32)
        btn.setObjectName("menuButton")
        return btn

    def _create_compound_widget(self):
        widget = QWidget()
        widget.setObjectName("menuItemContainer")
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(6, 0, 6, 0)
        layout.setSpacing(4)
        return widget

    def _create_divider(self):
        divider = QFrame()
        divider.setFrameShape(QFrame.VLine)
        divider.setFrameShadow(QFrame.Sunken)
        divider.setObjectName("divider")
        divider.setFixedWidth(1)
        return divider
        
    def show_menu(self):
        self.setWindowOpacity(0.0)
        self.show()
        self.fade_animation.setStartValue(0.0)
        self.fade_animation.setEndValue(1.0)
        self.fade_animation.start()

    def hide_menu(self):
        self.fade_animation.setStartValue(1.0)
        self.fade_animation.setEndValue(0.0)
        self.fade_animation.finished.connect(self.hide_after_animation)
        self.fade_animation.start()
    
    def hide_after_animation(self):
        self.hide()
        try:
            self.fade_animation.finished.disconnect(self.hide_after_animation)
        except TypeError:
            pass

# << CAMBIO: Nombre de la clase restaurado a TranscriptionWindow
class TranscriptionWindow(QWidget):
    # Señal para actualizar la UI desde otro hilo (Flask/transcripción)
    update_text_signal = pyqtSignal(dict)
    
    def __init__(self, model):
        super().__init__()
        # --- Integración de la lógica del backend ---
        self.model = model
        self.transcription_active = False
        self.start_time = 0
        self.transcriber_thread = None
        self.current_transcription_filepath = None
        self.current_audio_filepath = None
        self.selected_language = "es"
        self._already_stopped = False
        self.ngrok_tunnel = None
        self.public_url = None
        self.config_file = 'config.json'
        self.ia_settings_window = None
        self.main_menu_window = None
        self.flask_thread = threading.Thread(target=app.run, kwargs={'host': '0.0.0.0', 'port': 5000, 'debug': False, 'use_reloader': False})
        self.flask_thread.daemon = True
        self.flask_thread.start()
        
        # --- Configuración de la UI ---
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint | Qt.Tool)
        self.setAttribute(Qt.WA_TranslucentBackground)

        self._is_locked = False
        self.bg_color = QColor(25, 25, 25, int(0.85 * 255))
        self.font_color = QColor("#f5f5f5")
        self.current_font = QFont("Segoe UI", 14)
        
        self.offset = QPoint()
        self._resizing = False
        self._edge = 0
        self.initial_pos = QPoint()
        self.initial_geom = QRect()
        self.setMouseTracking(True)

        self.options_menu = OptionsMenu(self)
        self.options_menu.font_selector.setCurrentFont(self.current_font)
        
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)

        self.content_frame = QFrame(self)
        self.content_frame.setObjectName("contentFrame")
        self.content_frame.setMouseTracking(True)
        
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(25)
        shadow.setXOffset(0)
        shadow.setYOffset(8)
        shadow.setColor(QColor(0, 0, 0, 140))
        self.content_frame.setGraphicsEffect(shadow)

        content_layout = QVBoxLayout(self.content_frame)
        content_layout.setContentsMargins(5, 5, 5, 5)
        
        self.stacked_widget = QStackedWidget(self)
        self.stacked_widget.setMouseTracking(True)
        self.initial_screen = self._create_initial_screen()
        self.main_screen = self._create_main_screen()

        self.stacked_widget.addWidget(self.initial_screen)
        self.stacked_widget.addWidget(self.main_screen)
        
        content_layout.addWidget(self.stacked_widget)
        main_layout.addWidget(self.content_frame)
        
        self._load_styles()
        self._connect_signals()
        self._populate_devices()
        
        self.update_background_color()

        QTimer.singleShot(0, self.go_to_initial_screen)

    # =====================================================================
    # Métodos de Creación y Estilo de la UI
    # =====================================================================
    def _create_initial_screen(self):
        widget = QWidget()
        widget.setMouseTracking(True)
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(25, 8, 25, 8)
        layout.setSpacing(20)
        
        title_container = QWidget()
        title_layout = QHBoxLayout(title_container)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(10)
        mic_icon = QLabel()
        mic_icon.setPixmap(create_icon_from_svg(icon_data["mic"], "#0077b6").pixmap(24, 24))
        sttvar_label = QLabel("STTVAR")
        sttvar_label.setObjectName("welcomeTitle")
        title_layout.addWidget(mic_icon)
        title_layout.addWidget(sttvar_label)
        layout.addWidget(title_container)
        
        self.mic_selector = QComboBox()
        layout.addWidget(self.mic_selector)
        
        layout.addStretch(1)
        
        self.start_btn = QPushButton("Iniciar")
        self.start_btn.setIcon(create_icon_from_svg(icon_data["play"], "white"))
        self.start_btn.setObjectName("startButton")
        layout.addWidget(self.start_btn)

        self.initial_settings_btn = QPushButton()
        self.initial_settings_btn.setIcon(create_icon_from_svg(icon_data["settings"], "white"))
        self.initial_settings_btn.setObjectName("optionsToggleBtn")
        self.initial_settings_btn.setToolTip("Configurar Inteligencia Artificial")
        layout.addWidget(self.initial_settings_btn)

        self.initial_options_btn = QPushButton()
        self.initial_options_btn.setIcon(create_icon_from_svg(icon_data["files"], "white"))
        self.initial_options_btn.setObjectName("optionsToggleBtn")
        self.initial_options_btn.setToolTip("Abrir menú principal y herramientas")
        layout.addWidget(self.initial_options_btn)
        
        self.close_btn = QPushButton()
        self.close_btn.setIcon(create_icon_from_svg(icon_data["close"], "#a0a0a0"))
        self.close_btn.setObjectName("closeButton")
        self.close_btn.setToolTip("Cerrar aplicación")
        layout.addWidget(self.close_btn)
        
        return widget

    def _create_main_screen(self):
        widget = QWidget()
        widget.setMouseTracking(True)
        grid_layout = QGridLayout(widget)
        grid_layout.setContentsMargins(20, 4, 10, 4)
        
        self.transcript_text = QLabel("Iniciando transcripción...")
        self.transcript_text.setWordWrap(True)
        self.transcript_text.setAlignment(Qt.AlignCenter)
        self.transcript_text.setMouseTracking(True)
        self.transcript_text.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        
        self.options_btn = QPushButton()
        self.options_btn.setIcon(create_icon_from_svg(icon_data["menu"], "white"))
        self.options_btn.setObjectName("optionsToggleBtn")
        
        grid_layout.addWidget(self.transcript_text, 0, 0, 1, 1)
        grid_layout.addWidget(self.options_btn, 0, 1, Qt.AlignTop | Qt.AlignRight)
        
        return widget

    def _connect_signals(self):
        self.start_btn.clicked.connect(self.toggle_recording_state)
        self.close_btn.clicked.connect(self.close)
        self.initial_options_btn.clicked.connect(self._open_main_menu)
        self.initial_settings_btn.clicked.connect(self._open_ai_settings)
        self.options_btn.clicked.connect(self.toggle_options_menu)
        
        menu = self.options_menu
        menu.lock_btn.toggled.connect(self.toggle_lock)
        menu.rec_btn.toggled.connect(self.toggle_recording_state)
        menu.mute_btn.toggled.connect(self.toggle_mute)
        menu.share_btn.clicked.connect(self._on_qr_button_clicked)
        menu.bg_color_btn.clicked.connect(self.open_bg_color_picker)
        menu.font_color_btn.clicked.connect(self.open_font_color_picker)
        menu.opacity_slider.valueChanged.connect(self.change_opacity)
        menu.font_selector.currentFontChanged.connect(self.change_font_family)
        menu.language_selector.currentTextChanged.connect(self.change_language)
        menu.increase_font_btn.clicked.connect(lambda: self.change_font_size(1.1))
        menu.decrease_font_btn.clicked.connect(lambda: self.change_font_size(1/1.1))

        # Conectar la señal de actualización de texto al slot
        self.update_text_signal.connect(self._update_text_area)

    def _load_styles(self):
        self.setStyleSheet("""
            QToolTip {
                background-color: #2b2b2b; color: #f0f0f0; border: 1px solid #3c3c3c;
                padding: 5px; border-radius: 8px;
            }
            #optionsMenuBackground { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3a3a3a, stop:1 #282828);
                border: 1px solid #4a4a4a; border-radius: 19px; 
            }
            #menuButton, #menuItemContainer { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #4f4f4f, stop:1 #3e3e3e);
                border: 1px solid rgba(0,0,0, 0.2);
                border-top-color: rgba(255,255,255,0.05);
                border-radius: 12px; color: #f0f0f0;
            }
            #menuButton:hover, #menuItemContainer:hover { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5f5f5f, stop:1 #4c4c4c); 
            }
            #menuButton:pressed { background: #2a2a2a; }
            #menuButton:checked { background: #0077b6; border-color: #005a8e; }
            
            QPushButton#rec_btn, QPushButton#mic_btn, QPushButton#share_btn {
                background-color: #3c3c3c;
                border: 1px solid #555;
                border-radius: 16px;
            }
            QPushButton#rec_btn:hover, QPushButton#mic_btn:hover, QPushButton#share_btn:hover {
                background-color: #4d4d4d;
                border-color: #666;
            }
            QPushButton#rec_btn:pressed, QPushButton#mic_btn:pressed, QPushButton#share_btn:pressed {
                background-color: #2a2a2a;
            }
            
            QPushButton#rec_btn:checked, QPushButton#mic_btn:checked {
                background-color: #d90429;
                border: 1px solid #e63946;
            }
            QPushButton#rec_btn:checked:hover, QPushButton#mic_btn:checked:hover {
                 background-color: #e63946;
                 border-color: #f0505c;
            }

            #divider { background-color: rgba(255, 255, 255, 0.1); }
            #menuItemContainer QFontComboBox, #menuItemContainer QComboBox { 
                background: transparent; border: none; color: #f0f0f0; padding-left: 4px; 
            }
            #contentFrame { border-radius: 24px; color: #f5f5f5; }
            #welcomeTitle { font-size: 22px; font-weight: 700; background: transparent; color: #0077b6;}
            QComboBox { 
                padding: 4px 18px 4px 8px; border-radius: 10px; border: 1px solid rgba(255, 255, 255, 0.2); 
                background-color: rgba(40, 40, 40, 0.9); font-size: 13px; min-width: 120px;
                color: #f5f5f5;
            }
            QComboBox::drop-down { border: none; }
            #startButton { 
                padding: 8px 25px; font-size: 14px; font-weight: bold; color: white; 
                background-color: #0077b6; border: none; border-radius: 10px; 
            }
            #startButton:hover { background-color: #023e8a; }
            #closeButton, #optionsToggleBtn { background: transparent; border: none; padding: 8px; border-radius: 10px; }
            #closeButton:hover, #optionsToggleBtn:hover { background-color: rgba(255, 255, 255, 0.15); }
            QComboBox QAbstractItemView { 
                background-color: #282828; border: 1px solid #4a4a4a; 
                border-radius: 8px; color: #f5f5f5; selection-background-color: #0077b6; 
            }
            QComboBox QAbstractItemView::item { padding: 6px; }
            QSlider::groove:horizontal { height: 3px; background: rgba(255, 255, 255, 0.4); border-radius: 2px; }
            QSlider::handle:horizontal { 
                width: 14px; height: 14px; background: #f0f0f0; border-radius: 7px; 
                border: 1px solid rgba(0,0,0,0.2); margin: -5px 0; 
            }
        """)
        self.update_transcript_style()
        self.transcript_text.setFont(self.current_font)
        
    # =====================================================================
    # Métodos de Lógica de Transcripción y Backend
    # =====================================================================

    def _populate_devices(self):
        try:
            devices = sd.query_devices()
            input_devices = [(i, d['name']) for i, d in enumerate(devices) if d['max_input_channels'] > 0]
            if not input_devices:
                self.mic_selector.addItem("No hay micrófonos.")
                self.start_btn.setEnabled(False)
                return
            for idx, name in input_devices:
                self.mic_selector.addItem(name, idx)
        except Exception as e:
            QMessageBox.critical(self, "Error de Dispositivos", f"No se pudieron cargar los dispositivos de audio: {e}")

    def toggle_recording_state(self):
        # Este método es llamado por el botón de play y el de rec/stop
        is_checked = self.options_menu.rec_btn.isChecked()
        if self.transcription_active:
             if not is_checked: # Si la transcripción está activa y el botón se desmarca
                self._stop_transcription()
        else:
            if is_checked or not self.stacked_widget.currentWidget() == self.main_screen:
                self._start_transcription()

    def _start_transcription(self):
        global transcripcion_con_timestamps, transcripcion_plana, current_partial, qtextedit_buffer, shared_audio_data, transcription_chunks, shared_in_memory_audio_buffer
        
        device_index = self.mic_selector.currentData()
        if device_index is None:
            QMessageBox.warning(self, "Dispositivo no seleccionado", "Por favor, selecciona un dispositivo de micrófono.")
            self.options_menu.rec_btn.setChecked(False) # Revertir estado del botón
            return

        # Limpiar variables globales para una nueva sesión
        transcripcion_con_timestamps = ""
        transcripcion_plana = ""
        current_partial = ""
        qtextedit_buffer = ""
        transcription_chunks = []
        shared_in_memory_audio_buffer.clear()
        shared_audio_data.clear()

        self.start_time = time.time()
        
        try:
            engine = pyttsx3.init()
            engine.say("Iniciando grabación")
            engine.runAndWait()
        except Exception as e:
            print(f"No se pudo reproducir sonido de inicio: {e}")

        progress = QProgressDialog("Iniciando transcripción...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModal)
        progress.setCancelButton(None)
        progress.show()
        QApplication.processEvents()
        
        output_dir = "stt_guardados"
        audio_output_dir = "sttaudio_guardados"
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(audio_output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.current_transcription_filepath = os.path.join(output_dir, f"{timestamp}.txt")
        self.current_audio_filepath = os.path.join(audio_output_dir, f"{timestamp}.wav")

        try:
            self.transcriber_thread = TranscriberThread(self.model, device_index, self.current_transcription_filepath, self.current_audio_filepath, shared_in_memory_audio_buffer)
            self.transcriber_thread.new_text.connect(self.update_text_signal.emit)
            self.transcriber_thread.finished.connect(self._on_transcription_finished)
            self.transcriber_thread.start()
            
            self.transcription_active = True
            self._already_stopped = False
            self.transcriber_thread.set_mute(self.options_menu.mute_btn.isChecked())
            self.show_main_screen()
        except Exception as e:
            progress.close()
            QMessageBox.critical(self, "Error al Iniciar", str(e))
            self.transcription_active = False
            return
        finally:
            progress.close()

    def _stop_transcription(self):
        if self._already_stopped or not self.transcriber_thread:
            return
        self._already_stopped = True
        
        print("Deteniendo hilo de transcripción...")
        self.transcriber_thread.stop()
        self.transcriber_thread.wait()
        self.transcriber_thread = None
        
        self.transcription_active = False
        self.options_menu.rec_btn.setChecked(False)
        self.toggle_recording_visuals(False)
        
        self._prompt_save_or_discard()
        
        QTimer.singleShot(1500, self.go_to_initial_screen)

    def _on_transcription_finished(self):
        if self.transcription_active:
            print("El hilo de transcripción terminó inesperadamente.")
            self._stop_transcription()

    def _update_text_area(self, data: dict):
        global transcripcion_con_timestamps, transcripcion_plana, current_partial, qtextedit_buffer, last_qtextedit_update, transcription_chunks
        text_es = data.get("text", "")
        is_partial = data.get("is_partial", True)

        if text_es.strip() or not is_partial:
            current_time = time.time()
            if is_partial:
                qtextedit_buffer = text_es
                current_partial = text_es
            else:
                start_sec = data.get("start_sec")
                end_sec = data.get("end_sec")
                if start_sec is None: start_sec = time.time() - self.start_time
                timestamp_str = self._format_timestamp(start_sec)
                texto_limpio = text_es.strip()
                if texto_limpio:
                    chunk_data = {"timestamp": timestamp_str, "text": texto_limpio, "start_sec": start_sec, "end_sec": end_sec if end_sec is not None else start_sec + 2}
                    transcription_chunks.append(chunk_data)
                    transcripcion_con_timestamps += f"{timestamp_str} {texto_limpio}\n"
                    transcripcion_plana += f"{texto_limpio} "
                current_partial = ""
                qtextedit_buffer = texto_limpio

            if is_partial and (current_time - last_qtextedit_update < qtextedit_update_interval):
                return
            
            text_to_display = qtextedit_buffer
            if self.selected_language != "es" and text_to_display.strip() and source_lang:
                try:
                    target_lang = next((l for l in installed_languages if l.code == self.selected_language), None)
                    if target_lang:
                        text_to_display = source_lang.get_translation(target_lang).translate(text_to_display)
                except Exception as e:
                    print(f"Error al traducir en la ventana local: {e}")
                    text_to_display = "[Error de traducción] " + text_to_display
            
            self.transcript_text.setText(text_to_display)
            last_qtextedit_update = current_time

    def _format_timestamp(self, seconds: float) -> str:
        m, s = divmod(int(seconds), 60)
        h, m = divmod(m, 60)
        return f"[{h:02d}:{m:02d}:{s:02d}]"

    def _prompt_save_or_discard(self):
        global shared_audio_data
        has_transcription = self.current_transcription_filepath and os.path.exists(self.current_transcription_filepath)
        has_audio = self.current_audio_filepath and os.path.exists(self.current_audio_filepath)
        
        if not has_transcription and not has_audio:
            shared_audio_data.clear()
            self.transcript_text.setText("Sesión finalizada. No hay nada que guardar.")
            return

        self.transcript_text.setText("Procesando archivos finales...")
        QApplication.processEvents()

        save_transcription, save_audio, share_audio = False, False, False
        
        if has_transcription:
            reply = QMessageBox.question(self, "Guardar Transcripción", "¿Deseas guardar el archivo de texto de la transcripción?", QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            if reply == QMessageBox.Yes:
                save_transcription = True
        
        if has_audio:
            reply = QMessageBox.question(self, "Guardar Audio", "Guardar la grabación de audio es bajo tu propia responsabilidad.\n¿Deseas guardar el archivo de audio (.wav)?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes:
                save_audio = True
                reply_share = QMessageBox.question(self, "Compartir Audio", "¿Hacer este archivo de audio disponible para descargar en la sesión web?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                if reply_share == QMessageBox.Yes:
                    share_audio = True
                    shared_audio_data.update({
                        'directory': os.path.abspath(os.path.dirname(self.current_audio_filepath)),
                        'filename': os.path.basename(self.current_audio_filepath),
                        'available': True
                    })
        
        if not share_audio:
            shared_audio_data.clear()

        try:
            if not save_transcription and has_transcription:
                os.remove(self.current_transcription_filepath)
            if not save_audio and has_audio:
                os.remove(self.current_audio_filepath)
        except Exception as e:
            QMessageBox.warning(self, "Error al Eliminar", f"No se pudo eliminar un archivo temporal: {e}")

        messages = [msg for msg, flag in [
            ("Transcripción guardada.", save_transcription),
            ("Audio guardado.", save_audio),
            ("Audio compartido en la web.", share_audio)
        ] if flag]
        
        if messages:
            QMessageBox.information(self, "Operación Completada", "\n".join(messages))
        
        self.current_transcription_filepath = None
        self.current_audio_filepath = None
        self.transcript_text.setText("Listo para una nueva sesión.")

    # =====================================================================
    # Métodos de control de la UI y eventos
    # =====================================================================

    def center_on_screen(self):
        screen_geometry = QApplication.primaryScreen().geometry()
        x = (screen_geometry.width() - self.width()) // 2
        y = (screen_geometry.height() - self.height()) // 2
        self.move(x, y)

    def show_main_screen(self):
        if not self.transcription_active: return
        self.stacked_widget.setCurrentWidget(self.main_screen)
        self.setMinimumSize(400, 65)
        self.setMaximumSize(16777215, 16777215)
        self.resize(self.width(), 65)
        self.options_menu.rec_btn.setChecked(True)
        self.toggle_recording_visuals(True)
        
    def go_to_initial_screen(self):
        self.stacked_widget.setCurrentWidget(self.initial_screen)
        self.adjustSize()
        self.setFixedSize(self.size())
        self.options_menu.hide_menu()
        self.center_on_screen()
        
    def toggle_options_menu(self):
        if self.options_menu.isVisible():
            self.options_menu.hide_menu()
        else:
            frame_global_pos = self.content_frame.mapToGlobal(QPoint(0, 0))
            frame_width = self.content_frame.width()
            menu_size = self.options_menu.sizeHint()
            x = frame_global_pos.x() + (frame_width - menu_size.width()) / 2
            y = frame_global_pos.y() - menu_size.height() - 15
            self.options_menu.move(int(x), int(y))
            self.options_menu.show_menu()
            
    def change_language(self, lang_text):
        lang_map = {"Español": "es", "Inglés": "en", "Portugués": "pt"}
        self.selected_language = lang_map.get(lang_text, "es")
        print(f"Idioma de la UI local cambiado a: {self.selected_language}")
        if self.transcription_active and qtextedit_buffer.strip():
            self._update_text_area({"text": qtextedit_buffer, "is_partial": True})

    def toggle_lock(self, locked):
        self._is_locked = locked
        if locked:
            self.options_menu.lock_btn.setIcon(create_icon_from_svg(icon_data["lock"], "white"))
            self.setFixedSize(self.size())
        else:
            self.options_menu.lock_btn.setIcon(create_icon_from_svg(icon_data["unlock"], "white"))
            self.setMinimumSize(400, 65) 
            self.setMaximumSize(16777215, 16777215)

    def toggle_mute(self, muted):
        if self.transcriber_thread:
            self.transcriber_thread.set_mute(muted)
            
        menu = self.options_menu
        if muted:
            menu.mute_btn.setIcon(create_icon_from_svg(icon_data["mic_off"], "white"))
            menu.mute_btn.setToolTip("Reactivar micrófono")
        else:
            menu.mute_btn.setIcon(create_icon_from_svg(icon_data["mic"], "#f0f0f0"))
            menu.mute_btn.setToolTip("Silenciar micrófono")
            
    def toggle_recording_visuals(self, recording):
        menu = self.options_menu
        if recording:
            menu.rec_btn.setToolTip("Detener grabación")
            menu.rec_btn.setIcon(create_icon_from_svg(icon_data["stop"], "white"))
            self.transcript_text.setText("Escuchando...")
            menu.rec_animation.start()
        else:
            menu.rec_animation.stop()
            menu.rec_opacity_effect.setOpacity(1.0)
            menu.rec_btn.setToolTip("Iniciar grabación")
            menu.rec_btn.setIcon(create_icon_from_svg(icon_data["record_active"], "white"))
            self.transcript_text.setText("Grabación detenida. Procesando...")

    def open_bg_color_picker(self):
        color = QColorDialog.getColor(self.bg_color, self, "Seleccionar color de fondo")
        if color.isValid():
            self.bg_color.setRgb(color.red(), color.green(), color.blue(), self.bg_color.alpha())
            self.update_background_color()

    def open_font_color_picker(self):
        color = QColorDialog.getColor(self.font_color, self, "Seleccionar color de fuente")
        if color.isValid():
            self.font_color = color
            self.update_transcript_style()
    
    def change_opacity(self, value):
        self.bg_color.setAlpha(int(value * 2.55))
        self.update_background_color()

    def update_background_color(self):
        self.content_frame.setStyleSheet(f"""
            #contentFrame {{
                background-color: {self.bg_color.name(QColor.HexArgb)};
                border: 1px solid rgba(255, 255, 255, 0.2);
                border-radius: 24px;
            }}
        """)

    def change_font_family(self, font):
        self.current_font.setFamily(font.family())
        self.update_transcript_style()

    def change_font_size(self, factor):
        new_size = self.current_font.pointSizeF() * factor
        if new_size < 6: new_size = 6 
        self.current_font.setPointSizeF(new_size)
        self.update_transcript_style()

    def update_transcript_style(self):
        self.transcript_text.setFont(self.current_font)
        self.transcript_text.setStyleSheet(f"color: {self.font_color.name()}; background-color: transparent;")

    # =====================================================================
    # Métodos de Movimiento y Redimensión de la Ventana
    # =====================================================================
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton and not self._is_locked:
            self._edge = self._get_edge(event.pos())
            if self._edge != 0:
                self._resizing = True
                self.initial_pos = event.globalPos()
                self.initial_geom = self.geometry()
            elif self.content_frame.geometry().contains(event.pos()):
                self._resizing = False
                self.offset = event.globalPos() - self.pos()

    def mouseReleaseEvent(self, event):
        self._resizing = False
        self._edge = 0
        self.setCursor(Qt.ArrowCursor)

    def mouseMoveEvent(self, event):
        if not self._is_locked:
            if self._resizing:
                delta = event.globalPos() - self.initial_pos
                new_rect = QRect(self.initial_geom)
                
                if self._edge in [1, 7, 3]: new_rect.setLeft(self.initial_geom.left() + delta.x())
                if self._edge in [1, 5, 2]: new_rect.setTop(self.initial_geom.top() + delta.y())
                if self._edge in [2, 8, 4]: new_rect.setRight(self.initial_geom.right() + delta.x())
                if self._edge in [3, 6, 4]: new_rect.setBottom(self.initial_geom.bottom() + delta.y())
                
                if new_rect.width() < self.minimumSizeHint().width():
                    new_rect.setWidth(self.minimumSizeHint().width())
                if new_rect.height() < self.minimumSizeHint().height():
                    new_rect.setHeight(self.minimumSizeHint().height())

                self.setGeometry(new_rect)
            elif event.buttons() & Qt.LeftButton and not self._resizing:
                self.move(event.globalPos() - self.offset)
            else:
                current_edge = self._get_edge(event.pos())
                if current_edge != self._edge:
                    self._edge = current_edge
                    self._update_cursor_shape(self._edge)
    
    def _get_edge(self, pos):
        if self._is_locked or self.stacked_widget.currentWidget() != self.main_screen: return 0
        margin = 15
        rect = self.rect()
        on_left = pos.x() >= rect.left() and pos.x() < rect.left() + margin
        on_right = pos.x() > rect.right() - margin and pos.x() <= rect.right()
        on_top = pos.y() >= rect.top() and pos.y() < rect.top() + margin
        on_bottom = pos.y() > rect.bottom() - margin and pos.y() <= rect.bottom()
        
        if on_top and on_left: return 1;
        if on_top and on_right: return 2;
        if on_bottom and on_left: return 3;
        if on_bottom and on_right: return 4;
        if on_top: return 5;
        if on_bottom: return 6;
        if on_left: return 7;
        if on_right: return 8;
        return 0

    def _update_cursor_shape(self, edge):
        cursors = { 
            0: Qt.ArrowCursor, 1: Qt.SizeFDiagCursor, 2: Qt.SizeBDiagCursor, 
            3: Qt.SizeBDiagCursor, 4: Qt.SizeFDiagCursor, 5: Qt.SizeVerCursor, 
            6: Qt.SizeVerCursor, 7: Qt.SizeHorCursor, 8: Qt.SizeHorCursor 
        }
        self.setCursor(cursors.get(edge, Qt.ArrowCursor))
        
    # =====================================================================
    # Métodos de Navegación y Cierre
    # =====================================================================

    def _open_ai_settings(self):
        try:
            if not self.ia_settings_window or not self.ia_settings_window.isVisible():
                self.ia_settings_window = ConfiguracionIA(self, config_file=self.config_file)
                self.ia_settings_window.config_saved.connect(self._on_ai_config_saved)
                self.ia_settings_window.show()
        except NameError:
            QMessageBox.critical(self, "Error", "La ventana 'ConfiguracionIA' no está disponible. Revisa las importaciones.")

    def _on_ai_config_saved(self):
        print("Recargando configuración de IA tras guardado...")
        update_ollama_status()
        if not OLLAMA_AVAILABLE:
            QMessageBox.warning(self, "Advertencia de IA", OLLAMA_STATUS_MESSAGE)

    def _open_main_menu(self):
        try:
            if not self.main_menu_window or not self.main_menu_window.isVisible():
                self.main_menu_window = NuevaVentana(self)
                self.main_menu_window.show()
                # self.hide()  # <--- ELIMINA O COMENTA ESTA LÍNEA
        except NameError:
            QMessageBox.critical(self, "Error", "La ventana 'NuevaVentana' no está disponible. Revisa las importaciones.")

    def _get_local_ip(self):
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        try:
            s.connect(('10.255.255.255', 1))
            IP = s.getsockname()[0]
        except Exception: IP = '127.0.0.1'
        finally: s.close()
        return IP

    def _setup_ngrok_authtoken(self):
        token = None
        if os.path.exists(self.config_file):
            with open(self.config_file, 'r') as f:
                try: token = json.load(f).get('ngrok_authtoken')
                except json.JSONDecodeError: pass
        if not token:
            token, ok = QInputDialog.getText(self, 'Configuración de Ngrok', 'Por favor, introduce tu Authtoken de Ngrok.')
            if ok and token:
                config_data = {}
                if os.path.exists(self.config_file):
                    with open(self.config_file, 'r') as f:
                        config_data = json.load(f)
                config_data['ngrok_authtoken'] = token
                with open(self.config_file, 'w') as f: 
                    json.dump(config_data, f)
            else: return False
        try:
            ngrok.set_auth_token(token)
            return True
        except Exception as e:
            QMessageBox.critical(self, "Error de Ngrok", f"No se pudo configurar el token: {e}")
            return False

    def _on_qr_button_clicked(self):
        if self.ngrok_tunnel: self._show_qr_window(); return
        if not self._setup_ngrok_authtoken():
            QMessageBox.warning(self, "Cancelado", "Se necesita el Authtoken de Ngrok para compartir la sesión.")
            return
        progress = QProgressDialog("Creando enlace público...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModal); progress.setCancelButton(None); progress.show(); QApplication.processEvents()
        try:
            ngrok.kill()
            self.ngrok_tunnel = ngrok.connect(5000)
            self.public_url = self.ngrok_tunnel.public_url
            progress.close()
            self._show_qr_window()
        except Exception as e:
            progress.close()
            QMessageBox.critical(self, "Error de Ngrok", f"No se pudo crear el túnel: {e}")

    def _show_qr_window(self):
        local_ip = self._get_local_ip()
        public_url = self.public_url or "No disponible"
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            qr_app_path = os.path.join(current_dir, "qr_app.py")
            if not os.path.exists(qr_app_path):
                 QMessageBox.critical(self, "Error", f"No se encuentra el script 'qr_app.py' en la carpeta '{current_dir}'.")
                 return
            
            import subprocess
            subprocess.Popen([sys.executable, qr_app_path, f"http://{local_ip}:5000", public_url])
        except Exception as e:
            QMessageBox.critical(self, "Error al mostrar QR", f"No se pudo abrir la ventana de QR: {e}\n\nURL Local: http://{local_ip}:5000\nURL Pública: {public_url}")

    def closeEvent(self, event):
        print("Iniciando secuencia de cierre forzado...")

        # 1. Detener la transcripción si está activa (esto ya lo hacías bien)
        if self.transcription_active:
            print("Deteniendo hilo de transcripción...")
            self.transcriber_thread.stop()
            self.transcriber_thread.wait(2000) # Espera un máximo de 2 segundos

        # 2. Matar el túnel de Ngrok de forma explícita
        if self.ngrok_tunnel:
            print("Cerrando túnel de Ngrok...")
            try:
                ngrok.kill()
            except Exception as e:
                print(f"Error al intentar cerrar Ngrok (puede que ya estuviera cerrado): {e}")

        # 3. Forzar la terminación de todo el proceso de Python
        # Esta es la línea clave. os._exit(0) termina el programa inmediatamente,
        # cerrando todos los hilos (incluido Flask) sin contemplaciones.
        print("Terminando el proceso principal.")
        event.accept()  # Acepta el evento de cierre de la ventana
        os._exit(0)     # Fuerza la salida del script    # Fuerza la salida del script


if __name__ == '__main__':
    app = QApplication(sys.argv)
    
    try:
        from vosk import Model
        print("Cargando modelo de transcripción...")
        model_path = "model" 
        if not os.path.exists(model_path):
             QMessageBox.critical(None, "Error de Modelo", f"La carpeta del modelo de Vosk no se encontró en la ruta: '{model_path}'.\n\nDescarga un modelo, descomprímelo en una carpeta llamada 'model' y ponla junto al script.")
             sys.exit(1)
        model = Model(model_path)
        print("Modelo cargado exitosamente.")
    except Exception as e:
        QMessageBox.critical(None, "Error de Modelo", f"No se pudo cargar el modelo de Vosk. Asegúrate de que la biblioteca 'vosk' está instalada y la ruta al modelo es correcta.\n\nError: {e}")
        sys.exit(1)

    window = TranscriptionWindow(model=model)
    window.show()
    sys.exit(app.exec_())